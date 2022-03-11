from docx import Document
import re

from utils.Database import Database




class DocxUtils():
	def __init__(self, database, document) -> None: 
		self.doc = document
		self.dx = database

	def getFonts(self):
		fonts = set()
		for p in self.doc.paragraphs:
			fonts.add(p.style.font.size)
		return fonts
	
	def getFontOfEducation(self):
		for p in self.doc.paragraphs:
			if p.text.lower() == "education":
				return p.style.font.size
		return None

	def getHeaderFont(self):
		fonts = self.getFonts()
		educationFont = self.getFontOfEducation()
		for font in fonts:
			if font == educationFont:
				return font
		return None
	def getCapitalWords(self):
		headers = []
		for p in self.doc.paragraphs:
			if p.text.isupper():
				headers.append(p.text)
		return headers

	def getWords(self):
		words = []
		for p in self.doc.paragraphs:
			words += p.text.strip().split(' ')
		return words

	def searchDoc(self, listOfSearch:list, listOfParagraph:list, numRows:int):
		headers = []
		index = range(numRows)
		for indexOfPara, p in enumerate(listOfParagraph):
			if p.text == '':
				continue
			for i in index: 
				# header = list(filter(lambda e: isinstance(e, str) and e in p.text.lower() and len(p.text.split(' ')) < 4, listOfSearch[i]))
				try:
					# header = list(filter(lambda e: isinstance(e, str) and e == p.text.lower(), listOfSearch[i]))[0] 
					# headers.append([header,indexOfPara])
					headers += self.getHeader(listOfSearch[i], p.text, indexOfPara, headers)
				except IndexError:
					continue
		return headers
	
	def getHeader(self,listOfSearch:list, text:str, indexOfText:int, headers:list):
		_headers = []
		colNames = self.dx.getColumnsNames("points")
		for i, e in enumerate(listOfSearch):
			if isinstance(e, str) and e in text.lower().strip() and len(text.strip().split(' ')) < 3 and self.isHeaderExist(indexOfText, headers):
				_headers.append([colNames[i].strip(), indexOfText])
		return _headers
	
	def isHeaderExist(self, currentIndex:int, headers:list):
		for e in headers:
			if currentIndex == e[1]:
				return False
		return True
	
	def cleanText(self, text):
		return re.sub('[^A-Za-z0-9\@\-\.\,\(\)\[\]\"\'\:\#\*\+\%\ ]+', ' ', text)
	
	def getLink(self):
		# link = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
		links = []
		for key in self.doc.part.rels.keys():
			if self.doc.part.rels[key].is_external:
				links.append(self.doc.part.rels[key]._target)
		return links
	
	def getEmail(self, text):
		email = re.compile(r'[a-z0-9\.\-+_\:\/]+@[a-z0-9\.\-+_]+\.[a-z]+')
		return re.findall(email, text)

	def getPhone(self, text):
		phone = re.compile('\+?[0-9\ \-]{12,20}')
		return re.findall(phone, text)

	def getInfo(self, listOfParagraph:list):
		phones = []
		emails = []
		links = []
		for p in listOfParagraph:
			phone = self.getPhone(p.text)
			email = self.getEmail(p.text)
			# link = self.getLink(p.text)
			phones = phones+phone if len(phone)!=0 else phones
			emails = emails+email if len(email)!=0 else emails
			# links = links+link if len(link)!=0 else links
		links = self.getLink()
		return phones, emails, links
