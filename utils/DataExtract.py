from re import search
from docx import Document
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE

from utils.Database import Database
from utils.brain import TextProcesing

from .docxUtils import DocxUtils


class Extraction():
	def __init__(self, filename:str) -> None:
		self.filename = filename
		self.doc = Document(self.filename)
		self.dx = DocxUtils(self.filename)
		self.data = Database('Resume.db')
		self.process = TextProcesing()
		self.TABLENAME = "points"
		# self.QUERY = ["skills", "experience", "characteristics", "summary-objective", "education", "hobbies", "info"]
		pass

	def extractHeaders(self):
		fontHeader = self.dx.getHeaderFont()
		headers = []
		for p in self.doc.paragraphs:
			if p.style.font.size == fontHeader:
				headers.append(p.text)
		return headers
	def getQueries(self):
		query = []
		colNames = self.data.getColumnsNames(self.TABLENAME)
		colNames.pop(0)
		# for col in colNames: 
		# 	values = self.data.readColumn(self.TABLENAME, col)
		# 	q = ", ".join(values)
		# 	query.append(q)
		query = [
			"skills: Html, css",
			"experience, work history",
			"characteristics: team worker, Time management",
			"summary, objective, overview",
			"education: graduated college",
			"hobbies",
			"info, contact"
		]
		return colNames, query

	def extractHeadersV2(self):
		colNames, queries = self.getQueries()
		words = self.dx.getWords()
		sectionOfHeader = {}
		isHeader = False
		for word in words: 
			for index, q in enumerate(queries): 
				isHeader = self.process.sentence_similarity(word,q) 
				if isHeader: break
			if isHeader:
				sectionOfHeader[colNames[index]] = sectionOfHeader[colNames[index]]+" " + word if colNames[index] in sectionOfHeader else ""
			elif len(list(sectionOfHeader.keys())) > 0: 
				header = list(sectionOfHeader.keys())[-1] 
				sectionOfHeader[header] += " " + word
		return sectionOfHeader, sectionOfHeader[list(sectionOfHeader.keys())[0]]


	def extractSectionOfHeader(self):
		sectionOfHeaders = {}
		searches= self.data.read('points')
		headers = self.dx.searchDoc(listOfSearch=searches,listOfParagraph=self.doc.paragraphs, numRows=13)
		for i, header in enumerate(headers):
			sectionOfHeaders[header[0]] = [] if header[0] not in sectionOfHeaders else sectionOfHeaders[header[0]]
			if i == len(headers)-1:
				for index in range(header[1], len(self.doc.paragraphs)): 
					if header[0]!='info' and not self.process.cleanText(self.doc.paragraphs[index].text): 
						sectionOfHeaders[header[0]].append(self.doc.paragraphs[index].text.strip())
					if header[0]=="info": 
						sectionOfHeaders[header[0]] += [self.doc.paragraphs[index].text.strip()]
				continue
			for index in range(header[1]+1, headers[i+1][1]):
				if header[0]!='info' and not self.process.cleanText(self.doc.paragraphs[index].text): 
					sectionOfHeaders[header[0]] += [self.doc.paragraphs[index].text.strip()]
				if header[0] == 'info': 
					sectionOfHeaders[header[0]] += [self.doc.paragraphs[index].text.strip()]

		return sectionOfHeaders, headers[0] if len(headers)>0 else None
	
	def parseResume(self):

		#FIXME summary is hardcoded !!!!!!!!!!!!
		
		data, header = self.extractSectionOfHeader()
		if len(list(data.keys()))<2 or header == None: 
			data, header = self.extractHeadersV2()
		if 'info' not in data:
			phones, emails, links = self.dx.getInfo(self.doc.paragraphs)
			data['info'] = emails + links + phones
			data['info'] = list(filter(lambda x : x.strip()!="", data['info']))
		if 'summary' not in data:
			summList = self.process.summ(self.doc.paragraphs, header)
			data['summary'] = summList
		return data